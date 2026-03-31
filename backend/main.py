"""
OSB — Operating System Brain
FastAPI Backend Server
"""

from dotenv import load_dotenv
load_dotenv()

import os
import io
import re
import json
import traceback
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Optional
import uvicorn

app = FastAPI(title="OSB API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def get_groq_client():
    from groq import Groq
    groq_key = os.getenv("GROQ_API_KEY", "")
    if not groq_key:
        raise HTTPException(status_code=500, detail="GROQ_API_KEY not set")
    return Groq(api_key=groq_key)

def extract_text(raw: bytes, filename: str) -> str:
    fname = filename.lower()
    try:
        if fname.endswith(".pdf"):
            import fitz
            doc = fitz.open(stream=raw, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
            print(f"PDF extracted: {len(text)} chars")
            return text

        elif fname.endswith((".docx", ".doc")):
            import docx
            doc = docx.Document(io.BytesIO(raw))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        parts.append(row_text)
            text = "\n".join(parts)
            print(f"DOCX extracted: {len(text)} chars")
            return text

        elif fname.endswith((".xlsx", ".xls")):
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f"[Sheet: {sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    row_vals = [str(v) for v in row if v is not None and str(v).strip()]
                    if row_vals:
                        parts.append(" | ".join(row_vals))
            text = "\n".join(parts)
            print(f"Excel extracted: {len(text)} chars")
            return text

        elif fname.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp")):
            # Use Groq vision if available, otherwise describe
            try:
                import base64
                b64 = base64.b64encode(raw).decode()
                ext = fname.split('.')[-1]
                mime = f"image/{'jpeg' if ext in ['jpg','jpeg'] else ext}"
                return f"__IMAGE_B64__{mime}__SEP__{b64}"
            except:
                return f"[Image file: {filename}]"

        elif fname.endswith(".csv"):
            text = raw.decode("utf-8", errors="ignore")
            lines = text.split('\n')
            # Take header + all rows but limit total
            return "\n".join(lines[:200])

        else:
            return raw.decode("utf-8", errors="ignore")

    except Exception as ex:
        print(f"Extraction error for {filename}: {ex}")
        try:
            return raw.decode("utf-8", errors="ignore")
        except:
            return f"[Could not read file: {filename}]"


@app.get("/")
def root():
    return {"message": "OSB API is running", "status": "operational"}

@app.get("/health")
def health():
    key = os.getenv("GROQ_API_KEY", "")
    return {"status": "ok", "groq_key_set": bool(key)}

@app.post("/query")
async def query(question: str = Form(...), history: str = Form(default="[]")):
    """General query with no files"""
    print(f"General query: {question}")
    try:
        client = get_groq_client()
        messages = [{"role": "system", "content": "You are OSB — Operating System Brain, an intelligent AI assistant. Answer thoroughly and helpfully."}]

        # Add conversation history
        try:
            hist = json.loads(history)
            for h in hist[-3:]:
                messages.append({"role": "user", "content": h["q"]})
                messages.append({"role": "assistant", "content": h["a"]})
        except: pass

        messages.append({"role": "user", "content": question})

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            max_tokens=1200,
            messages=messages
        )
        answer = response.choices[0].message.content
        return {"answer": answer, "sources": []}
    except HTTPException:
        raise
    except Exception as e:
        print("EXCEPTION:", traceback.format_exc())
        raise HTTPException(status_code=500, detail=type(e).__name__ + ": " + str(e))


@app.post("/query-with-files")
async def query_with_files(
    question: str = Form(...),
    files: List[UploadFile] = File(default=[]),
    history: str = Form(default="[]")
):
    print(f"Received question: {question}")
    print(f"Received files: {[f.filename for f in files]}")
    try:
        client = get_groq_client()

        # Build messages with history
        messages = []
        try:
            hist = json.loads(history)
            for h in hist[-3:]:
                messages.append({"role": "user", "content": h["q"]})
                messages.append({"role": "assistant", "content": h["a"]})
        except: pass

        # Process files
        file_context = ""
        image_files = []

        for f in files:
            raw = await f.read()
            text = extract_text(raw, f.filename)

            if text.startswith("__IMAGE_B64__"):
                # Handle image with vision
                parts = text.split("__SEP__")
                mime = parts[0].replace("__IMAGE_B64__", "")
                b64 = parts[1]
                image_files.append({"filename": f.filename, "mime": mime, "b64": b64})
            else:
                file_context += f"\n--- FILE: {f.filename} ---\n{text[:8000]}\n--- END ---\n"

        # Build system prompt
        system_prompt = """You are OSB — Operating System Brain, an intelligent document analysis AI.

PART 1 — FROM THE FILE:
Extract and quote relevant information directly from the provided file content.
Clearly state what the file says about the topic. Mention the filename as source.

PART 2 — ADDITIONAL CONTEXT:
Expand on the same topic using your broader knowledge.
Add real-world context and details that complement the file."""

        if file_context:
            system_prompt += f"\n\nFILE CONTENTS:\n{file_context}"

        # Handle images with vision model
        if image_files:
            user_content = [{"type": "text", "text": question}]
            for img in image_files:
                user_content.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:{img['mime']};base64,{img['b64']}"}
                })
            messages_with_sys = [{"role": "system", "content": system_prompt}] + messages + [{"role": "user", "content": user_content}]
            try:
                response = client.chat.completions.create(
                    model="llama-3.2-11b-vision-preview",
                    max_tokens=1200,
                    messages=messages_with_sys
                )
            except Exception:
                # Fallback to text model with image description
                fallback_msgs = [{"role": "system", "content": system_prompt}] + messages + [{"role": "user", "content": f"{question}\n[Image file uploaded: {', '.join(img['filename'] for img in image_files)}]"}]
                response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    max_tokens=1200,
                    messages=fallback_msgs
                )
        else:
            messages_with_sys = [{"role": "system", "content": system_prompt}] + messages + [{"role": "user", "content": question}]
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                max_tokens=1200,
                messages=messages_with_sys
            )

        answer = response.choices[0].message.content
        sources = [f.filename for f in files]
        print(f"Answer generated: {len(answer)} chars")
        return {"answer": answer, "sources": sources}

    except HTTPException:
        raise
    except Exception as e:
        print("EXCEPTION:", traceback.format_exc())
        raise HTTPException(status_code=500, detail=type(e).__name__ + ": " + str(e))


if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
